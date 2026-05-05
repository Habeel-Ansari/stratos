from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "stratos-energy-website-content.md"
OUTPUT = BASE_DIR / "Stratos Energy Website Content Draft.docx"


def add_bold_label_paragraph(document: Document, line: str) -> None:
    paragraph = document.add_paragraph()
    if ":  " in line:
        label, value = line.split(":  ", 1)
    elif ": " in line:
        label, value = line.split(": ", 1)
    else:
        paragraph.add_run(line)
        return
    paragraph.add_run(f"{label}:").bold = True
    paragraph.add_run(f" {value}")


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    if "—" in text:
        raise ValueError("Source content contains an em dash.")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(22)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(16)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3.font.size = Pt(11)

    lines = text.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = document.add_paragraph(stripped[2:], style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("## "):
            document.add_paragraph(stripped[3:], style="Heading 1")
            continue

        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:], style="Heading 2")
            continue

        if stripped.startswith("**") and stripped.endswith("**"):
            document.add_paragraph(stripped[2:-2], style="Heading 3")
            continue

        if stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. "):
            document.add_paragraph(stripped, style="List Number")
            continue

        if stripped.endswith("  "):
            stripped = stripped[:-2]

        if stripped.startswith("Low ") or stripped.startswith("Medium ") or stripped.startswith("High ") or stripped.startswith("Power ") or stripped.startswith("Control ") or stripped.startswith("Instrumentation ") or stripped.startswith("Project-specific ") or stripped.startswith("Cable trays") or stripped.startswith("Cable ladders") or stripped.startswith("Conduits and fittings") or stripped.startswith("Cable glands") or stripped.startswith("Clamps and supports") or stripped.startswith("Routing accessories") or stripped.startswith("Conduit fittings") or stripped.startswith("Connectors") or stripped.startswith("Couplings") or stripped.startswith("Clamps and straps") or stripped.startswith("Reducers and adapters") or stripped.startswith("Bushings") or stripped.startswith("Nipples") or stripped.startswith("Conduit bodies") or stripped.startswith("Sealing fittings") or stripped.startswith("Hardware and accessories") or stripped.startswith("Grounding busbars") or stripped.startswith("Neutral bars") or stripped.startswith("Compression connectors") or stripped.startswith("Mechanical connectors and clamps") or stripped.startswith("Exothermic materials") or stripped.startswith("Bare conductors") or stripped.startswith("Air rods") or stripped.startswith("Strike pads") or stripped.startswith("Bonding clamps") or stripped.startswith("Lightning protection accessories") or stripped.startswith("Junction boxes") or stripped.startswith("Terminal enclosures") or stripped.startswith("Weather-resistant enclosures") or stripped.startswith("Industrial enclosures") or stripped.startswith("Accessories for indoor and outdoor applications") or stripped.startswith("Miniature circuit breakers") or stripped.startswith("Moulded case circuit breakers") or stripped.startswith("Air circuit breakers") or stripped.startswith("Contactors") or stripped.startswith("Safety switches") or stripped.startswith("Distribution boards") or stripped.startswith("Control components") or stripped.startswith("Disconnector switches") or stripped.startswith("Industrial plugs and sockets") or stripped.startswith("Surge protection devices") or stripped.startswith("OPGW cable") or stripped.startswith("Telecom cable") or stripped.startswith("Fiber-related cable solutions") or stripped.startswith("Utility and infrastructure cable solutions") or stripped.startswith("Application-specific cable supply") or stripped.startswith("Cable transit systems") or stripped.startswith("Pipe sealing systems") or stripped.startswith("Modular sealing blocks") or stripped.startswith("Frames and accessories") or stripped.startswith("Multi-cable and pipe entry sealing solutions") or stripped.startswith("Industrial lighting") or stripped.startswith("Commercial lighting") or stripped.startswith("Indoor lighting") or stripped.startswith("Outdoor lighting") or stripped.startswith("Area lighting") or stripped.startswith("Roadway lighting") or stripped.startswith("Architectural lighting") or stripped.startswith("Utility lighting accessories") or stripped.startswith("Network accessories") or stripped.startswith("Communication cables") or stripped.startswith("Routers") or stripped.startswith("Switches") or stripped.startswith("Gateways") or stripped.startswith("Structured communication installation materials") or stripped in {"Oil and Gas", "Construction", "Telecommunication", "Marine", "Energy", "Manufacturing", "Mining", "Petrochemical", "Infrastructure and Utilities"}:
            document.add_paragraph(stripped, style="List Bullet")
            continue

        if stripped.startswith("Company Name:") or stripped.startswith("Website:") or stripped.startswith("Primary Email:") or stripped.startswith("Sales Email:") or stripped.startswith("Phone:") or stripped.startswith("Address:") or stripped.startswith("Prepared for:") or stripped.startswith("Reference structure used:") or stripped.startswith("Important note:") or stripped.startswith("Quick Links:") or stripped.startswith("Contact:") or stripped.startswith("Sales:") or stripped.startswith("Meta Title:") or stripped.startswith("Meta Description:"):
            add_bold_label_paragraph(document, stripped)
            continue

        document.add_paragraph(stripped)

    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
